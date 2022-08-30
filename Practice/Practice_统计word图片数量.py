# -*- coding: UTF-8 -*-

"""
Author: Ethan Chai
Date: 2022/8/29 20:42

input: 

output: 

Short Description: 
    - 知识点：
        - 参考：https://hellowac.github.io/docx_doc/tutorial.html

Change History:
    
"""
from docx import Document  # 导入docx包
from docx.shared import Pt, Cm

document = Document()
document.add_paragraph('111111111')  # 添加第1个段落
document.add_paragraph('2222222222')  # 添加第2个段落
document.add_paragraph('3333333333')  # 添加第3个段落
document.add_picture(r'../EthanFileData/Images/3ac96999b99f41f7bcd47414c40ac1dd.png')  # 文档末尾添加名称为111.jpg的图像
document.add_paragraph('111111111')  # 添加第5个段落
document.add_paragraph('2222222222')  # 添加第6个段落
document.add_paragraph('3333333333')  # 添加第7个段落
document.add_picture(r'../EthanFileData/Images/3ac96999b99f41f7bcd47414c40ac1dd.png')  # 文档末尾添加名称为111.jpg的图像
print('段落数量：', len(document.paragraphs))
document.save('test.docx')  # 保存docx文档

document1 = Document('test.docx')
print('图形图像的数量：', len(document1.inline_shapes))
for paragraph in document.paragraphs:
    print(paragraph.text)
